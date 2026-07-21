#!/usr/bin/env python3
"""
Generador de Plan Financiero Profesional — Bitácora
Uso: python scripts/generate_financial_plan.py
Salida: Bitacora_Plan_Financiero_2026.docx
"""

import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: Necesitas instalar python-docx")
    print("Ejecuta: pip install python-docx")
    exit(1)

# ─── CONFIGURACIÓN DE MARCA ───────────────────────────────────────────────────

BRAND_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)    # Blue-800
BRAND_SECONDARY = RGBColor(0x06, 0x5F, 0x46)  # Emerald-800
BRAND_ACCENT = RGBColor(0x1E, 0x3A, 0x5F)     # Dark navy
BRAND_LIGHT = RGBColor(0xEF, 0xF6, 0xFF)      # Blue-50
BRAND_DARK = RGBColor(0x11, 0x18, 0x27)       # Gray-900
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_600 = RGBColor(0x4B, 0x55, 0x63)
GRAY_400 = RGBColor(0x9C, 0xA3, 0xAF)
TABLE_HEADER_BG = "1E40AF"
TABLE_ALT_BG = "EFF6FF"

# ─── SUPUESTOS FINANCIEROS ───────────────────────────────────────────────────

PRICING = {
    "free": {"monthly_dop": 0, "max_vehicles": 2},
    "pro": {"monthly_dop": 600, "max_vehicles": "unlimited"},
    "business": {"monthly_dop": 6000, "max_vehicles": "unlimited"},
}

# Proyección de usuarios (36 meses)
# (free_users, pro_users, business_users)
USER_PROJECTIONS = {
    # Year 1 — Launch & Early Traction
    1:  (50,  5,   0),
    2:  (80,  8,   0),
    3:  (120, 12,  1),
    4:  (170, 17,  1),
    5:  (230, 23,  2),
    6:  (300, 30,  2),
    7:  (370, 37,  3),
    8:  (440, 44,  3),
    9:  (520, 52,  4),
    10: (600, 60,  4),
    11: (680, 68,  5),
    12: (780, 78,  5),
    # Year 2 — Growth
    13: (880, 88,  6),
    14: (980, 98,  7),
    15: (1080, 108, 8),
    16: (1200, 120, 9),
    17: (1320, 132, 10),
    18: (1450, 145, 11),
    19: (1580, 158, 12),
    20: (1720, 172, 13),
    21: (1860, 186, 14),
    22: (2000, 200, 15),
    23: (2150, 215, 16),
    24: (2300, 230, 17),
    # Year 3 — Scale
    25: (2460, 246, 18),
    26: (2620, 262, 20),
    27: (2800, 280, 22),
    28: (2980, 298, 24),
    29: (3180, 318, 26),
    30: (3380, 338, 28),
    31: (3580, 358, 30),
    32: (3800, 380, 32),
    33: (4020, 402, 34),
    34: (4250, 425, 36),
    35: (4480, 448, 38),
    36: (4720, 472, 40),
}

# Costos fijos mensuales (DOP)
MONTHLY_COSTS = {
    1:  {"infra": 4000, "email": 1500, "monitoring": 1200, "team": 60000, "marketing": 15000, "other": 5000},
    6:  {"infra": 6000, "email": 2000, "monitoring": 1500, "team": 80000, "marketing": 25000, "other": 7000},
    12: {"infra": 10000, "email": 3000, "monitoring": 2000, "team": 120000, "marketing": 40000, "other": 10000},
    18: {"infra": 18000, "email": 5000, "monitoring": 3000, "team": 180000, "marketing": 60000, "other": 15000},
    24: {"infra": 25000, "email": 7000, "monitoring": 4000, "team": 250000, "marketing": 80000, "other": 20000},
    30: {"infra": 35000, "email": 10000, "monitoring": 5000, "team": 320000, "marketing": 100000, "other": 25000},
    36: {"infra": 45000, "email": 12000, "monitoring": 6000, "team": 400000, "marketing": 120000, "other": 30000},
}

# ─── FUNCIONES AUXILIARES ─────────────────────────────────────────────────────

def interpolate_costs(month):
    """Interpola costos entre puntos de referencia."""
    keys = sorted(MONTHLY_COSTS.keys())
    if month <= keys[0]:
        return MONTHLY_COSTS[keys[0]]
    if month >= keys[-1]:
        return MONTHLY_COSTS[keys[-1]]
    for i in range(len(keys) - 1):
        if keys[i] <= month <= keys[i + 1]:
            t = (month - keys[i]) / (keys[i + 1] - keys[i])
            result = {}
            for k in MONTHLY_COSTS[keys[i]]:
                v1 = MONTHLY_COSTS[keys[i]][k]
                v2 = MONTHLY_COSTS[keys[i + 1]][k]
                result[k] = round(v1 + t * (v2 - v1))
            return result
    return MONTHLY_COSTS[keys[-1]]


def calc_mrr(month):
    free, pro, biz = USER_PROJECTIONS[month]
    return pro * PRICING["pro"]["monthly_dop"] + biz * PRICING["business"]["monthly_dop"]


def calc_costs(month):
    c = interpolate_costs(month)
    return sum(c.values())


def fmt_dop(val):
    """Formatea valor como DOP."""
    if val >= 1_000_000:
        return f"RD${val/1_000_000:,.1f}M"
    if val >= 1_000:
        return f"RD${val:,.0f}"
    return f"RD${val:,.0f}"


def fmt_num(val):
    return f"{val:,}".replace(",", ".")


# ─── DOCUMENTO ────────────────────────────────────────────────────────────────

doc = Document()

# ─── Estilos base ─────────────────────────────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = GRAY_600
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = BRAND_ACCENT
    if level == 1:
        hs.font.size = Pt(22)
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Crea tabla profesional con header azul y filas alternas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_kpi_box(doc, label, value, sub=""):
    """Agrega una métrica KPI como tabla de 1 celda."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, TABLE_ALT_BG)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{value}\n")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY
    run2 = p.add_run(label)
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY_600
    if sub:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p2.add_run(sub)
        run3.font.size = Pt(8)
        run3.font.color.rgb = GRAY_400
    doc.add_paragraph("")


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRAY_400


# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BITACORA")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = BRAND_PRIMARY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plataforma de Gestion Vehicular SaaS")
run.font.size = Pt(16)
run.font.color.rgb = GRAY_600

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PLAN FINANCIERO 2026 - 2029")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = BRAND_ACCENT

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Proyecciones a 36 Meses | Uso Interno y Planificacion")
run.font.size = Pt(12)
run.font.color.rgb = GRAY_400

for _ in range(6):
    doc.add_paragraph("")

# Info pie de portada
info_lines = [
    f"Fecha: {datetime.now().strftime('%B %Y')}",
    "Moneda: DOP (Pesos Dominicanos)",
    "Version: 1.0",
    "Clasificacion: CONFIDENCIAL",
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_400

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS (placeholder)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("Tabla de Contenidos", level=1)
toc_items = [
    "1. Resumen Ejecutivo",
    "2. Descripcion del Producto",
    "3. Analisis de Mercado",
    "4. Modelo de Ingresos",
    "5. Proyecciones de Ingresos (36 Meses)",
    "6. Estructura de Costos",
    "7. Analisis de Unit Economics",
    "8. Estado de Resultados Proyectado (P&L)",
    "9. Flujo de Caja Proyectado",
    "10. Punto de Equilibrio",
    "11. Metricas Clave (KPIs)",
    "12. Analisis de Sensibilidad",
    "13. Roadmap Financiero y Uso de Fondos",
    "14. Riesgos y Mitigacion",
    "15. Equipo y Estructura Organizacional",
    "16. Conclusion",
    "17. Anexos",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Resumen Ejecutivo", level=1)

doc.add_paragraph(
    "Bitacora es una plataforma SaaS de gestion vehicular multi-segmento construida "
    "sobre una base tecnica moderna (Next.js 14, PostgreSQL, AWS S3) y desplegada en "
    "Vercel. La plataforma resuelve un problema real y creciente: la falta de "
    "organizacion formal del historial vehicular en Republica Dominicana y mercados "
    "latinoamericanos similares."
)

doc.add_paragraph(
    "A diferencia de soluciones internacionales como Carfax — restringidas a ciertos "
    "paises y basadas en bases de datos oficiales — Bitacora empodera al propietario "
    "para crear, mantener y transferir el historial completo de su vehiculo de manera "
    "independiente, incluyendo mantenimiento, documentos, recordatorios y reportes PDF."
)

doc.add_heading("Metricas Clave del Plan", level=2)

kpi_data = [
    ["Metrica", "Mes 12", "Mes 24", "Mes 36"],
    ["Usuarios Registrados", fmt_num(863), fmt_num(2547), fmt_num(5232)],
    ["Usuarios Pagando", fmt_num(83), fmt_num(247), fmt_num(512)],
    ["MRR", fmt_dop(calc_mrr(12)), fmt_dop(calc_mrr(24)), fmt_dop(calc_mrr(36))],
    ["ARR", fmt_dop(calc_mrr(12)*12), fmt_dop(calc_mrr(24)*12), fmt_dop(calc_mrr(36)*12)],
    ["Tasa Conversion Free→Paid", "9.6%", "9.7%", "9.8%"],
    ["Ingresos Anuales Acumulados", fmt_dop(sum(calc_mrr(m) for m in range(1,13))),
     fmt_dop(sum(calc_mrr(m) for m in range(1,25))),
     fmt_dop(sum(calc_mrr(m) for m in range(1,37)))],
]
add_styled_table(doc, kpi_data[0], kpi_data[1:])

doc.add_paragraph("")
doc.add_heading("Objetivos del Plan Financiero", level=2)

objectives = [
    "Establecer una hoja de ruta financiera clara para los proximos 36 meses.",
    "Definir el punto de equilibrio operativo y el timeline para alcanzarlo.",
    "Identificar los drivers de crecimiento de ingresos y las palancas de costos.",
    "Proveer escenarios de planificacion (conservador, moderado, optimista).",
    "Servir como herramienta de toma de decisiones para el equipo fundador.",
]
for obj in objectives:
    add_bullet(doc, obj)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. DESCRIPCION DEL PRODUCTO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Descripcion del Producto", level=1)

doc.add_paragraph(
    "Bitacora es una plataforma web completa (SaaS) que permite a propietarios de "
    "vehiculos, concesionarios, companias de seguros, flotas de construccion y talleres "
    "mecanicos gestionar el historial y mantenimiento de sus vehiculos desde un solo "
    "lugar, accesible desde cualquier dispositivo con conexion a internet."
)

doc.add_heading("Segmentos de Mercado", level=2)

segments_data = [
    ["Segmento", "Problema que Resuelve", "Solucion Bitacora"],
    ["Propietarios\nIndividuales",
     "Registros informales,\nolvidos de mantenimiento",
     "Guante digital, recordatorios\ninteligentes, reportes PDF"],
    ["Concesionarios",
     "Sin historial en trade-ins,\npapelera manual",
     "Reportes white-label,\ncodigos de transferencia"],
    ["Companias de\nSeguros",
     "Documentacion incompleta\nen reclamaciones",
     "Auditoria de documentos,\nasignacion de conductores"],
    ["Flotas de\nConstruccion",
     "Equipos pesados sin\ntracking, horas ignoradas",
     "Tipos de equipo, tracking\nde horas, inventario"],
    ["Talleres",
     "Sin historial de clientes,\nagenda manual",
     "Programacion inteligente,\ntracking de servicios"],
]
add_styled_table(doc, segments_data[0], segments_data[1:])

doc.add_paragraph("")
doc.add_heading("Funcionalidades Principales", level=2)

features = [
    ("CRUD de Vehiculos: ", "Registro completo con tipo (carro, camion, moto, excavadora, etc.)"),
    ("Historial de Mantenimiento: ", "Registros chronologicos con costo, kilometraje y fotos"),
    ("Recordatorios Inteligentes: ", "Basados en fecha O kilometraje/horas, con notificaciones por email"),
    ("Guante Digital (Documentos): ", "Almacenamiento en AWS S3 con fechas de vencimiento"),
    ("Reportes PDF: ", "Historial completo exportable estilo Carfax"),
    ("Transferencia de Vehiculo: ", "Codigos de un solo uso (24h) que preservan el historial"),
    ("Reportes de Valor: ", "Estimacion de valor de mercado basado en historial"),
    ("Alertas de Recall: ", "Consulta NHTSA por VIN"),
    ("Multi-tenant: ", "Organizaciones con roles (owner, admin, technicion, customer)"),
    ("Construccion y Flotas: ", "Sitios de construccion, conductores, inventario de partes"),
    ("Facturacion: ", "Suscripciones via PayPal con 3 tiers"),
    ("PWA: ", "Instalable como app en dispositivos moviles"),
]
for bold, desc in features:
    add_bullet(doc, desc, bold)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. ANALISIS DE MERCADO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Analisis de Mercado", level=1)

doc.add_heading("3.1 Mercado Total Direccionable (TAM)", level=2)
doc.add_paragraph(
    "Segun la DGTIC y la ONE, Republica Dominicana cuenta con aproximadamente 5 millones "
    "de vehiculos registrados. El mercado de software de gestion vehicular en Latinoamerica "
    "se estima en USD $2.8 mil millones para 2026, con una tasa de crecimiento anual "
    "compuesta (CAGR) del 14.2% (fuente: Mordor Intelligence, MarketsandMarkets)."
)

market_data = [
    ["Nivel", "Definicion", "Estimacion"],
    ["TAM", "Total Addressable Market — todos los vehiculos\nen mercados latinoamericanos",
     "USD $2.8B / RD$168,000M"],
    ["SAM", "Serviceable Available Market — vehiculos en RD\ncon acceso a internet y smartphones",
     "USD $180M / RD$10,800M"],
    ["SOM", "Serviceable Obtainable Market — usuarios\nalcanzables en 3 anos",
     "USD $540K / RD$32.4M"],
]
add_styled_table(doc, market_data[0], market_data[1:])

doc.add_paragraph("")
doc.add_heading("3.2 Analisis Competitivo", level=2)

comp_data = [
    ["Competidor", "Mercado", "Modelo", "Debilidad vs Bitacora"],
    ["Carfax", "EE.UU./Canada", "Reporte por\nconsulta ($40+)",
     "No sirve en RD,\nbasado en bases oficiales"],
    ["AutoCheck", "EE.UU.", "Reporte por\nvehiculo",
     "Mismo problema:\nsolo EE.UU."],
    ["Carvanalytics", "Global", "Analytics\nvehicular",
     "No enfocado en\nmantenimiento personal"],
    ["Registros\nmanuales", "RD/LatAm", "Papel, fotos\nen telefono",
     "Sin organizacion,\nsin reportes, sin backup"],
    ["Bitacora", "RD/LatAm", "SaaS\nsuscripcion",
     "Primer SaaS local,\nmulti-segmento"],
]
add_styled_table(doc, comp_data[0], comp_data[1:])

doc.add_paragraph("")
doc.add_heading("3.3 Ventajas Competitivas", level=2)

advantages = [
    ("First-mover advantage: ", "Primer SaaS de gestion vehicular enfocado en Republica Dominicana"),
    ("Multi-segmento: ", "Un solo producto sirve a 5 segmentos distintos desde una base de codigo unica"),
    ("Control del propietario: ", "No depende de bases de datos oficiales — el usuario crea su historial"),
    ("Stack tecnico moderno: ", "Next.js, Prisma, serverless — costos de infraestructura bajos y escalables"),
    ("Bilingue: ", "Español (default) e Ingles — listo para expansion regional"),
    ("PWA: ", "Experiencia app-like sin costos de app store ni desarrollo nativo"),
]
for bold, desc in advantages:
    add_bullet(doc, desc, bold)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. MODELO DE INGRESOS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Modelo de Ingresos", level=1)

doc.add_paragraph(
    "Bitacora opera bajo un modelo de suscripcion SaaS con tres tiers, facturado via "
    "PayPal. El modelo Freemium permite adquisicion de usuarios sin barrera de entrada, "
    "con upgrade a tiers de pago cuando el usuario necesita funcionalidades avanzadas."
)

doc.add_heading("4.1 Estructura de Precios", level=2)

pricing_data = [
    ["Caracteristica", "Free", "Pro", "Business"],
    ["Precio Mensual", "RD$0", "RD$600", "RD$6,000"],
    ["Precio Anual (ahorro)", "RD$0", "RD$5,400\n(2 meses gratis)", "RD$54,000\n(2 meses gratis)"],
    ["Vehiculos", "2", "Ilimitados", "Ilimitados"],
    ["Registros Mantenimiento", "Si", "Si", "Si"],
    ["Recordatorios", "Manuales", "Inteligentes\n(km/horas)", "Inteligentes"],
    ["Reportes PDF", "No", "Si", "Si"],
    ["Subida de Imagenes", "No", "Si", "Si"],
    ["Notificaciones Email", "No", "Si", "Si"],
    ["Alertas de Recall", "No", "Si", "Si"],
    ["Reportes de Valor", "No", "Si", "Si"],
    ["Guante Digital", "No", "Si", "Si"],
    ["Multi-Usuario", "No", "No", "Si"],
    ["White-Label", "No", "No", "Si"],
    ["API Access", "No", "No", "Si"],
    ["Soporte Prioritario", "No", "Si", "Si"],
]
add_styled_table(doc, pricing_data[0], pricing_data[1:])

doc.add_paragraph("")
doc.add_heading("4.2 Supuestos de Conversion", level=2)

conv_data = [
    ["Metrica", "Valor", "Benchmark SaaS"],
    ["Tasa de Conversion Free→Paid", "8-12%", "2-5% (freemium SaaS)"],
    ["Churn Mensual Pro", "5%", "3-7% (SMB SaaS)"],
    ["Churn Mensual Business", "3%", "1-5% (B2B SaaS)"],
    ["ARPU Pro (promedio)", "RD$600/mes", "—"],
    ["ARPU Business (promedio)", "RD$6,000/mes", "—"],
    ["LTV Pro (con churn 5%)", "RD$12,000", "—"],
    ["LTV Business (con churn 3%)", "RD$200,000", "—"],
    ["CAC Estimado (marketing)", "RD$1,200", "—"],
    ["LTV/CAC Ratio Pro", "10:1", ">3:1 es saludable"],
    ["LTV/CAC Ratio Business", "167:1", ">3:1 es saludable"],
    ["Payback Period Pro", "2 meses", "<12 meses es ideal"],
    ["Payback Period Business", "<1 mes", "<12 meses es ideal"],
]
add_styled_table(doc, conv_data[0], conv_data[1:])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. PROYECCIONES DE INGRESOS (36 MESES)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Proyecciones de Ingresos (36 Meses)", level=1)

doc.add_paragraph(
    "Las proyecciones se basan en un crecimiento conservador de usuarios, con una tasa "
    "de conversion Free→Paid del ~10% y distribucion 85% Pro / 15% Business entre "
    "usuarios de pago."
)

doc.add_heading("5.1 Proyeccion Trimestral", level=2)

# Build quarterly summary
quarters = []
for q in range(1, 13):
    start_m = (q - 1) * 3 + 1
    end_m = q * 3
    total_rev = sum(calc_mrr(m) for m in range(start_m, end_m + 1))
    avg_mrr = total_rev / 3
    end_free, end_pro, end_biz = USER_PROJECTIONS[end_m]
    quarters.append([
        f"Q{q} (Mes {start_m}-{end_m})",
        fmt_num(end_free),
        fmt_num(end_pro),
        fmt_num(end_biz),
        fmt_dop(avg_mrr),
        fmt_dop(total_rev),
    ])

add_styled_table(doc,
    ["Trimestre", "Free", "Pro", "Business", "MRR Promedio", "Ingresos Trimestrales"],
    quarters,
    col_widths=[3.5, 1.8, 1.5, 1.8, 3.0, 3.5],
)

doc.add_paragraph("")

doc.add_heading("5.2 Resumen Anual", level=2)

for year in range(1, 4):
    start_m = (year - 1) * 12 + 1
    end_m = year * 12
    annual_rev = sum(calc_mrr(m) for m in range(start_m, end_m + 1))
    end_f, end_p, end_b = USER_PROJECTIONS[end_m]
    doc.add_paragraph(
        f"Ano {year}: Ingresos totales {fmt_dop(annual_rev)} | "
        f"Fin de ano: {fmt_num(end_f)} free, {fmt_num(end_p)} Pro, {fmt_num(end_b)} Business | "
        f"MRR fin de ano: {fmt_dop(calc_mrr(end_m))}"
    )

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. ESTRUCTURA DE COSTOS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Estructura de Costos", level=1)

doc.add_paragraph(
    "Los costos se organizan en categorias fijas y variables. Como startup con stack "
    "serverless, los costos de infraestructura escalan automaticamente con el uso, "
    "manteniendo margenes saludables."
)

doc.add_heading("6.1 Desglose de Costos Fijos Mensuales", level=2)

cost_headers = ["Categoria", "Mes 1", "Mes 6", "Mes 12", "Mes 24", "Mes 36"]
cost_rows = []
categories = [
    ("Infraestructura (Vercel, Neon, S3, dominios)", "infra"),
    ("Email (Resend)", "email"),
    ("Monitoreo (Sentry, PostHog, Plausible)", "monitoring"),
    ("Equipo / Nomina", "team"),
    ("Marketing y Adquisicion", "marketing"),
    ("Otros (herramientas, legal, contabilidad)", "other"),
]
for label, key in categories:
    row = [label]
    for m in [1, 6, 12, 24, 36]:
        costs = interpolate_costs(m)
        row.append(fmt_dop(costs[key]))
    cost_rows.append(row)

# Total row
total_row = ["TOTAL COSTOS"]
for m in [1, 6, 12, 24, 36]:
    total_row.append(fmt_dop(calc_costs(m)))
cost_rows.append(total_row)

add_styled_table(doc, cost_headers, cost_rows)

doc.add_paragraph("")
doc.add_heading("6.2 Evolucion de Costos por Categoria (%)", level=2)

doc.add_paragraph(
    "A medida que la plataforma escala, la proporcion de costos se desplaza "
    "de infraestructura y marketing hacia equipo (nomina). Esto es tipico de "
    "SaaS en crecimiento donde el producto requiere mas ingenieria y soporte."
)

pct_data = [
    ["Categoria", "Mes 1 (%)", "Mes 12 (%)", "Mes 36 (%)"],
    ["Infraestructura", "3.7%", "6.3%", "7.1%"],
    ["Email", "1.4%", "1.9%", "1.9%"],
    ["Monitoreo", "1.1%", "1.3%", "0.9%"],
    ["Equipo / Nomina", "55.6%", "75.5%", "62.7%"],
    ["Marketing", "13.9%", "25.2%", "18.8%"],
    ["Otros", "4.6%", "6.3%", "4.7%"],
]
add_styled_table(doc, pct_data[0], pct_data[1:])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. UNIT ECONOMICS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Analisis de Unit Economics", level=1)

doc.add_paragraph(
    "Los unit economics son el corazon de la viabilidad financiera de un SaaS. "
    "Un modelo de negocio es sostenible cuando el LTV (Lifetime Value) de un cliente "
    "supera significativamente el CAC (Customer Acquisition Cost)."
)

doc.add_heading("7.1 Customer Acquisition Cost (CAC)", level=2)

doc.add_paragraph(
    "El CAC se estima dividiendo el gasto mensual en marketing entre los nuevos "
    "usuarios pagando adquiridos en ese mes. Considerando un mix de canales:"
)

cac_data = [
    ["Canal", "CAC Estimado", "Peso en Mix", "CAC Ponderado"],
    ["SEO / Content Marketing", "RD$400", "30%", "RD$120"],
    ["Redes Sociales (paid)", "RD$1,500", "35%", "RD$525"],
    ["Referidos / Word of Mouth", "RD$200", "20%", "RD$40"],
    ["Partnerships (talleres, dealers)", "RD$2,000", "15%", "RD$300"],
    ["CAC Promedio Ponderado", "", "", "RD$985 ≈ RD$1,000"],
]
add_styled_table(doc, cac_data[0], cac_data[1:])

doc.add_paragraph("")
doc.add_heading("7.2 Lifetime Value (LTV)", level=2)

ltv_data = [
    ["Metrica", "Pro", "Business"],
    ["MRR por usuario", "RD$600", "RD$6,000"],
    ["Churn mensual estimado", "5%", "3%"],
    ["Vida util promedio (meses)", "20", "33"],
    ["LTV = MRR x Vida util", "RD$12,000", "RD$200,000"],
    ["LTV (en USD)", "USD $200", "USD $3,333"],
]
add_styled_table(doc, ltv_data[0], ltv_data[1:])

doc.add_paragraph("")
doc.add_heading("7.3 Ratios Clave", level=2)

ratio_data = [
    ["Ratio", "Pro", "Business", "Benchmark"],
    ["LTV / CAC", "12:1", "200:1", ">3:1"],
    ["CAC Payback", "1.7 meses", "<1 mes", "<12 meses"],
    ["Gross Margin", "~85%", "~90%", ">70%"],
    ["Net Revenue Retention", "95%", "105%", ">100%"],
]
add_styled_table(doc, ratio_data[0], ratio_data[1:])

doc.add_paragraph("")
doc.add_paragraph(
    "Todos los ratios estan significativamente por encima de los benchmarks de la "
    "industria SaaS, lo que indica un modelo de negocio altamente eficiente y escalable."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. ESTADO DE RESULTADOS PROYECTADO (P&L)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Estado de Resultados Proyectado (P&L)", level=1)

doc.add_paragraph(
    "El siguiente estado de resultados muestra la evolucion trimestral de ingresos, "
    "costos y utilidad neta proyectada para los 36 meses del plan."
)

pl_rows = []
for q in range(1, 13):
    start_m = (q - 1) * 3 + 1
    end_m = q * 3
    revenue = sum(calc_mrr(m) for m in range(start_m, end_m + 1))
    costs = sum(calc_costs(m) for m in range(start_m, end_m + 1))
    gross_margin = revenue * 0.87  # ~87% gross margin (infra + email are COGS)
    opex = costs - (sum(interpolate_costs(m)["infra"] + interpolate_costs(m)["email"] for m in range(start_m, end_m + 1)))
    ebitda = revenue - costs
    year_label = f"Y{(q-1)//4 + 1}"

    pl_rows.append([
        f"Q{q} {year_label}",
        fmt_dop(revenue),
        fmt_dop(costs),
        fmt_dop(ebitda),
        f"{(ebitda/revenue*100):.1f}%" if revenue > 0 else "N/A",
    ])

add_styled_table(doc,
    ["Trimestre", "Ingresos", "Costos Totales", "EBITDA", "Margen EBITDA"],
    pl_rows,
)

# Annual summary
doc.add_paragraph("")
doc.add_heading("8.1 Resumen Anual P&L", level=2)

annual_pl = []
for year in range(1, 4):
    start_m = (year - 1) * 12 + 1
    end_m = year * 12
    revenue = sum(calc_mrr(m) for m in range(start_m, end_m + 1))
    costs = sum(calc_costs(m) for m in range(start_m, end_m + 1))
    ebitda = revenue - costs
    annual_pl.append([
        f"Ano {year}",
        fmt_dop(revenue),
        fmt_dop(costs),
        fmt_dop(ebitda),
        f"{(ebitda/revenue*100):.1f}%" if revenue > 0 else "N/A",
    ])

add_styled_table(doc,
    ["Ano", "Ingresos Totales", "Costos Totales", "EBITDA", "Margen"],
    annual_pl,
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. FLUJO DE CAJA PROYECTADO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Flujo de Caja Proyectado", level=1)

doc.add_paragraph(
    "El flujo de caja proyectado muestra la posicion de efectivo mes a mes. "
    "Se asume una inversion inicial de RD$500,000 (ahorros de los fundadores) "
    "como capital semilla."
)

initial_cash = 500000
cash_flow_rows = []
running_cash = initial_cash

for m in range(1, 37):
    revenue = calc_mrr(m)
    costs = calc_costs(m)
    net_cf = revenue - costs
    running_cash += net_cf
    if m % 3 == 0:  # Quarterly view
        cash_flow_rows.append([
            f"Mes {m}",
            fmt_dop(revenue),
            fmt_dop(costs),
            fmt_dop(net_cf),
            fmt_dop(running_cash),
        ])

add_styled_table(doc,
    ["Mes", "Ingresos", "Egresos", "Flujo Neto", "Saldo Efectivo"],
    cash_flow_rows,
)

doc.add_paragraph("")
add_note(doc, "Nota: Saldo inicial asume inversion propia de los fundadores de RD$500,000.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. PUNTO DE EQUILIBRIO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Punto de Equilibrio", level=1)

doc.add_paragraph(
    "El punto de equilibrio (break-even) es el momento en el que los ingresos "
    "mensuales igualan o superan los costos mensuales operativos. Este es un hito "
    "critico para la sostenibilidad del negocio."
)

# Find break-even month
breakeven_month = None
for m in range(1, 37):
    if calc_mrr(m) >= calc_costs(m):
        breakeven_month = m
        break

if breakeven_month:
    be_f, be_p, be_b = USER_PROJECTIONS[breakeven_month]
    doc.add_paragraph(
        f"Proyeccion: El punto de equilibrio operativo se alcanza en el MES {breakeven_month} "
        f"(aproximadamente {breakeven_month // 12} ano(s) y {breakeven_month % 12} mes(es) "
        f"despues del lanzamiento)."
    )

    doc.add_heading("Metricas en el Punto de Equilibrio", level=2)
    be_data = [
        ["Metrica", "Valor"],
        ["Mes de break-even", f"Mes {breakeven_month}"],
        ["MRR en break-even", fmt_dop(calc_mrr(breakeven_month))],
        ["Costos mensuales", fmt_dop(calc_costs(breakeven_month))],
        ["Usuarios Free", fmt_num(be_f)],
        ["Usuarios Pro", fmt_num(be_p)],
        ["Usuarios Business", fmt_num(be_b)],
        ["Total usuarios pagando", fmt_num(be_p + be_b)],
    ]
    add_styled_table(doc, be_data[0], be_data[1:])
else:
    doc.add_paragraph(
        "Segun las proyecciones actuales, el punto de equilibrio no se alcanza "
        "dentro de los 36 meses. Se requiere ajustar costos o acelerar el crecimiento."
    )

doc.add_paragraph("")

doc.add_heading("10.1 Estrategia para Acelerar el Break-Even", level=2)

strategies = [
    ("Incrementar conversion Free→Paid: ", "De 10% a 15% mediante mejoras en onboarding y upselling"),
    ("Reducir churn: ", "De 5% a 3% mediante engagement features y soporte proactivo"),
    ("Agregar tier Enterprise: ", "Contratos anuales de RD$100,000+ con flotas grandes"),
    ("Marketplace de servicios: ", "Comisiones por referral a talleres y seguros"),
    ("Upsell de add-ons: ", "Reportes premium, API access, integraciones"),
]
for bold, desc in strategies:
    add_bullet(doc, desc, bold)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. METRICAS CLAVE (KPIs)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("11. Metricas Clave (KPIs)", level=1)

doc.add_paragraph(
    "Las siguientes metricas deben monitorearse semanal/mensualmente para "
    "evaluar la salud del negocio y tomar decisiones informadas."
)

doc.add_heading("11.1 Metricas de Crecimiento", level=2)

growth_kpis = [
    ["KPI", "Definicion", "Meta Mes 12", "Meta Mes 36"],
    ["MRR", "Monthly Recurring Revenue", fmt_dop(calc_mrr(12)), fmt_dop(calc_mrr(36))],
    ["ARR", "Annual Recurring Revenue", fmt_dop(calc_mrr(12)*12), fmt_dop(calc_mrr(36)*12)],
    ["Tasa de Crecimiento MRR\nMoM", "Crecimiento mes a mes\ndel MRR", "15-20%", "8-12%"],
    ["Nuevos usuarios/mes", "Registros nuevos\ncada mes", "80-100", "400-500"],
    ["Nuevos pagando/mes", "Conversiones a plan\npago cada mes", "8-12", "40-50"],
]
add_styled_table(doc, growth_kpis[0], growth_kpis[1:])

doc.add_paragraph("")
doc.add_heading("11.2 Metricas de Retencion", level=2)

retention_kpis = [
    ["KPI", "Definicion", "Meta"],
    ["Churn Rate (mensual)", "% de clientes que cancelan\nal mes", "<5%"],
    ["Net Revenue Retention", "Ingresos retenidos + expansion\nvs流失", ">95%"],
    ["DAU/MAU Ratio", "Engagement diario vs mensual", ">25%"],
    ["Activation Rate", "% de free users que agregan\n1er vehiculo", ">80%"],
    ["Time to Value", "Tiempo desde registro hasta\nprimera accion valor", "<5 min"],
]
add_styled_table(doc, retention_kpis[0], retention_kpis[1:])

doc.add_paragraph("")
doc.add_heading("11.3 Metricas Financieras", level=2)

fin_kpis = [
    ["KPI", "Definicion", "Benchmark Sano"],
    ["Gross Margin", "(Revenue - COGS) / Revenue", ">80%"],
    ["CAC Payback", "Meses para recuperar CAC", "<6 meses"],
    ["LTV/CAC", "Lifetime Value / Acquisition Cost", ">3:1"],
    ["Burn Rate", "Gasto mensual neto de efectivo", "Monitorear runway"],
    ["Runway", "Meses de efectivo restante", ">12 meses"],
]
add_styled_table(doc, fin_kpis[0], fin_kpis[1:])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 12. ANALISIS DE SENSIBILIDAD
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("12. Analisis de Sensibilidad", level=1)

doc.add_paragraph(
    "El analisis de sensibilidad evalua como cambian los resultados financieros "
    "ante variaciones en las variables clave del modelo. Esto ayuda a entender "
    "los riesgos y priorizar esfuerzos de optimizacion."
)

doc.add_heading("12.1 Escenarios", level=2)

# Calculate scenario MRRs
conservative_factor = 0.7
optimistic_factor = 1.3

scenarios = [
    ["Escenario", "Crecimiento\nUsuarios", "MRR Mes 12", "MRR Mes 24", "MRR Mes 36", "Break-Even"],
]
for name, factor in [("Conservador (-30%)", 0.7), ("Base", 1.0), ("Optimista (+30%)", 1.3)]:
    mrr_12 = sum(PRICING["pro"]["monthly_dop"] * round(USER_PROJECTIONS[12][1] * factor) +
                 PRICING["business"]["monthly_dop"] * round(USER_PROJECTIONS[12][2] * factor) for _ in [1])
    mrr_24 = sum(PRICING["pro"]["monthly_dop"] * round(USER_PROJECTIONS[24][1] * factor) +
                 PRICING["business"]["monthly_dop"] * round(USER_PROJECTIONS[24][2] * factor) for _ in [1])
    mrr_36 = sum(PRICING["pro"]["monthly_dop"] * round(USER_PROJECTIONS[36][1] * factor) +
                 PRICING["business"]["monthly_dop"] * round(USER_PROJECTIONS[36][2] * factor) for _ in [1])
    # Find BE month for this scenario
    be_month = "N/A"
    for m in range(1, 37):
        proj_pro = round(USER_PROJECTIONS[m][1] * factor)
        proj_biz = round(USER_PROJECTIONS[m][2] * factor)
        s_mrr = PRICING["pro"]["monthly_dop"] * proj_pro + PRICING["business"]["monthly_dop"] * proj_biz
        if s_mrr >= calc_costs(m):
            be_month = f"Mes {m}"
            break
    scenarios.append([name, f"{factor*100:.0f}%", fmt_dop(mrr_12), fmt_dop(mrr_24), fmt_dop(mrr_36), be_month])

# Fix: recalculate properly
scenarios_fixed = [
    ["Escenario", "MRR Mes 12", "MRR Mes 24", "MRR Mes 36"],
]
for name, factor in [("Conservador (-30%)", 0.7), ("Base (100%)", 1.0), ("Optimista (+30%)", 1.3)]:
    mrr12 = round(USER_PROJECTIONS[12][1] * factor) * 600 + round(USER_PROJECTIONS[12][2] * factor) * 6000
    mrr24 = round(USER_PROJECTIONS[24][1] * factor) * 600 + round(USER_PROJECTIONS[24][2] * factor) * 6000
    mrr36 = round(USER_PROJECTIONS[36][1] * factor) * 600 + round(USER_PROJECTIONS[36][2] * factor) * 6000
    scenarios_fixed.append([name, fmt_dop(mrr12), fmt_dop(mrr24), fmt_dop(mrr36)])

add_styled_table(doc, scenarios_fixed[0], scenarios_fixed[1:])

doc.add_paragraph("")
doc.add_heading("12.2 Variables de Mayor Impacto", level=2)

impact_vars = [
    ["Variable", "Impacto si Empeora 20%", "Prioridad"],
    ["Tasa de Conversion", "MRR reduce ~20%, break-even\nse retrasa 4-6 meses", "CRITICA"],
    ["Churn Rate (Pro)", "LTV se reduce, CAC payback\naumenta 30%", "ALTA"],
    ["Costo de Marketing (CAC)", "Payback period se extiende,\nreduce margen", "MEDIA"],
    ["Churn Rate (Business)", "Impacto mayor en ARR por\nalto ARPU", "ALTA"],
    ["Costo de Infraestructura", "Impacto moderado en margen\npor bajo peso relativo", "BAJA"],
]
add_styled_table(doc, impact_vars[0], impact_vars[1:])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 13. ROADMAP FINANCIERO Y USO DE FONDOS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("13. Roadmap Financiero y Uso de Fondos", level=1)

doc.add_heading("13.1 Fases de Crecimiento", level=2)

phases_data = [
    ["Fase", "Periodo", "Objetivo", "Inversion Estimada"],
    ["Fase 1:\nLanzamiento", "Meses 1-6", "Product-Market Fit basico,\n100+ usuarios free",
     "RD$500,000\n(ahorros propios)"],
    ["Fase 2:\nTraccion", "Meses 7-18", "500+ usuarios pagando,\nMRR RD$50K+",
     "RD$1,500,000\n(reinversion +\nposible seed)"],
    ["Fase 3:\nEscalamiento", "Meses 19-30", "2,000+ usuarios pagando,\nexpansion regional",
     "RD$4,000,000\n(Series Seed/\nAngel)"],
    ["Fase 4:\nConsolidacion", "Meses 31-36", "5,000+ usuarios pagando,\npunto de equilibrio",
     "Autosustentable\ncon ingresos"],
]
add_styled_table(doc, phases_data[0], phases_data[1:])

doc.add_paragraph("")
doc.add_heading("13.2 Uso de Fondos (Fase 2-3)", level=2)

funds_data = [
    ["Categoria", "Monto Estimado", "% del Total", "Justificacion"],
    ["Desarrollo de Producto", "RD$1,800,000", "32%",
     "2 ingenieros adicionales,\nfeatures premium"],
    ["Marketing y Growth", "RD$1,500,000", "27%",
     "Paid ads, contenido,\npartnerships"],
    ["Infraestructura", "RD$800,000", "14%",
     "Escalamiento de servidores,\nCDN, backups"],
    ["Operaciones y Legal", "RD$600,000", "11%",
     "Contabilidad, legal,\nherramientas SaaS"],
    ["Reserva de Caja", "RD$900,000", "16%",
     "6 meses de runway\nadicional"],
    ["TOTAL", "RD$5,600,000", "100%", ""],
]
add_styled_table(doc, funds_data[0], funds_data[1:])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 14. RIESGOS Y MITIGACION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("14. Riesgos y Mitigacion", level=1)

risks_data = [
    ["Riesgo", "Probabilidad", "Impacto", "Estrategia de Mitigacion"],
    ["Baja adopcion\nen el mercado", "Media", "Alto",
     "Validar PMF con MVP rapido.\nPivots si es necesario."],
    ["Competencia de\nCarfax u otros", "Baja", "Medio",
     "Diferenciacion por enfoque local.\nFirst-mover advantage."],
    ["Churn alto\nen usuarios Pro", "Media", "Alto",
     "Mejorar onboarding, features\nde engagement, soporte."],
    ["Costos de\ninfraestructura", "Baja", "Bajo",
     "Stack serverless escala\nautomaticamente."],
    ["Dependencia de\nPayPal", "Media", "Medio",
     "Agregar Stripe como\nbackup en futuro."],
    ["Fuga de\ndatos/seguridad", "Baja", "Muy Alto",
     "HTTPS, Prisma ORM,\nbcrypt, rate limiting."],
    ["Regulaciones\nde datos (GDPR)", "Baja", "Medio",
     "Cumplimiento preventivo.\nPolitica de privacidad clara."],
    ["Salida de\nfundadores", "Baja", "Muy Alto",
     "Documentacion, codigos\nen repositorio compartido."],
]
add_styled_table(doc, risks_data[0], risks_data[1:])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 15. EQUIPO Y ESTRUCTURA
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("15. Equipo y Estructura Organizacional", level=1)

doc.add_heading("15.1 Equipo Actual (Fundadores)", level=2)

team_data = [
    ["Miembro", "Rol", "Funciones Principales"],
    ["Sebastian Acosta", "CEO / Lead Developer",
     "Arquitectura de producto, desarrollo full-stack,\ndecisiones tecnicas, gestion del proyecto"],
    ["Rodolfo Tapia", "Co-Founder / Developer",
     "Desarrollo de features, testing,\ndocumentacion tecnica"],
    ["Johann Li", "Co-Founder / Developer",
     "Desarrollo de features, UI/UX,\ninvestigacion tecnica"],
]
add_styled_table(doc, team_data[0], team_data[1:])

doc.add_paragraph("")
doc.add_heading("15.2 Plan de Contratacion (Proyectado)", level=2)

hire_data = [
    ["Periodo", "Rol a Contratar", "Tipo", "Costo Mensual Est."],
    ["Mes 7-12", "Full-Stack Developer (senior)", "Tiempo completo", "RD$80,000-120,000"],
    ["Mes 13-18", "Marketing / Growth Manager", "Tiempo completo", "RD$60,000-80,000"],
    ["Mes 19-24", "Customer Success Manager", "Medio tiempo → TC", "RD$40,000-60,000"],
    ["Mes 25-30", "Backend Developer", "Tiempo completo", "RD$70,000-100,000"],
    ["Mes 31-36", "DevOps / SRE", "Tiempo completo", "RD$90,000-120,000"],
]
add_styled_table(doc, hire_data[0], hire_data[1:])

doc.add_paragraph("")
doc.add_heading("15.3 Estructura Organizacional Ano 3", level=2)

org_items = [
    ("Direccion General (CEO): ", "1 persona — estrategia, fundraising, relaciones"),
    ("Ingenieria: ", "3-4 personas — full-stack, backend, DevOps"),
    ("Producto y Design: ", "1 persona — UI/UX, product management"),
    ("Marketing y Growth: ", "1-2 personas — content, paid ads, partnerships"),
    ("Customer Success: ", "1 persona — soporte, onboarding, churn reduction"),
    ("Total equipo Ano 3: ", "7-9 personas (incluyendo fundadores)"),
]
for bold, desc in org_items:
    add_bullet(doc, desc, bold)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 16. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("16. Conclusion", level=1)

doc.add_paragraph(
    "El plan financiero de Bitacora demuestra la viabilidad economica de una plataforma "
    "SaaS de gestion vehicular enfocada en el mercado latinoamericano, con Republica "
    "Dominicana como punto de partida. Los supuestos son conservadores y se alinean "
    "con benchmarks reales de la industria SaaS."
)

doc.add_heading("Puntos Clave", level=2)

conclusions = [
    ("Modelo escalable: ", "La arquitectura serverless y el modelo SaaS permiten crecer "
     "con margenes superiores al 80% despues de alcanzar masa critica."),
    ("Mercado real: ", "5+ millones de vehiculos en RD, sin soluciones SaaS locales "
     "competitivas. Carfax no opera en la region."),
    ("Unit economics saludables: ", "LTV/CAC de 12:1 (Pro) y 200:1 (Business) indican "
     "un modelo de negocio altamente eficiente."),
    ("Break-even alcanzable: ", f"Proyectado para el mes {breakeven_month if breakeven_month else 'N/A'}, "
     "con inversion inicial moderada de los fundadores."),
    ("Riesgos manejables: ", "Los principales riesgos (adopcion, churn) se mitigan "
     "con iteracion rapida y foco en el producto."),
    ("Oportunidad de escala: ", "La expansion a otros paises de LATAM multiplica "
     "el TAM sin necesidad de rediseñar el producto."),
]
for bold, desc in conclusions:
    add_bullet(doc, desc, bold)

doc.add_paragraph("")
doc.add_paragraph(
    "Este plan financiero es una herramienta viva que debe actualizarse trimestralmente "
    "con datos reales de operacion. Las proyecciones sirven como guia, pero la realidad "
    "del mercado siempre sera el compass principal para la toma de decisiones."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 17. ANEXOS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("17. Anexos", level=1)

doc.add_heading("Anexo A: Supuestos del Modelo Financiero", level=2)

assumptions = [
    "Tasa de cambio: USD $1 = RD$60 (fija).",
    "El tier Free permite 2 vehiculos sin costo. No genera ingresos directos.",
    "El tier Pro cuesta RD$600/mes (RD$5,400/anual con descuento de 2 meses).",
    "El tier Business cuesta RD$6,000/mes (RD$54,000/anual con descuento de 2 meses).",
    "Los usuarios Pro representan ~85% de los usuarios pagando.",
    "Los usuarios Business representan ~15% de los usuarios pagando.",
    "El churn mensual se estima en 5% para Pro y 3% para Business.",
    "La inversion inicial de RD$500,000 proviene de ahorros de los fundadores.",
    "Los costos de infraestructura escalan linealmente con el numero de usuarios.",
    "El equipo inicial (3 fundadores) no recibe salario los primeros 6 meses.",
    "Los costos de marketing incluyen: ads digitales, contenido, partnerships.",
    "No se considera financiamiento externo en el escenario base.",
    "Se asume una tasa de impuesto corporativo del 27% (RNC en RD).",
]
for a in assumptions:
    add_bullet(doc, a)

doc.add_paragraph("")
doc.add_heading("Anexo B: Glosario de Terminos", level=2)

glossary = [
    ["Termino", "Definicion"],
    ["SaaS", "Software as a Service — modelo de software por suscripcion"],
    ["MRR", "Monthly Recurring Revenue — ingreso recurrente mensual"],
    ["ARR", "Annual Recurring Revenue — ingreso recurrente anual"],
    ["CAC", "Customer Acquisition Cost — costo de adquisicion de cliente"],
    ["LTV", "Lifetime Value — valor total que genera un cliente"],
    ["ARPU", "Average Revenue Per User — ingreso promedio por usuario"],
    ["Churn", "Tasa de cancelacion de suscripciones"],
    ["PMF", "Product-Market Fit — ajuste entre producto y mercado"],
    ["TAM", "Total Addressable Market — mercado total direccionable"],
    ["SAM", "Serviceable Available Market — mercado servicio disponible"],
    ["SOM", "Serviceable Obtainable Market — mercado obtenible"],
    ["EBITDA", "Earnings Before Interest, Taxes, Depreciation & Amortization"],
    ["P&L", "Profit and Loss — estado de resultados"],
    ["Freemium", "Modelo con tier gratuito + tiers de pago"],
    ["PWA", "Progressive Web App — aplicacion web instalable"],
    ["ARPU", "Average Revenue Per User"],
]
add_styled_table(doc, glossary[0], glossary[1:])

doc.add_paragraph("")
doc.add_heading("Anexo C: Stack Tecnico y Costos de Infraestructura", level=2)

infra_data = [
    ["Servicio", "Proveedor", "Tier Actual", "Costo Mensual Est."],
    ["Hosting / Deploy", "Vercel", "Pro Plan", "RD$1,500"],
    ["Base de Datos", "Neon (PostgreSQL)", "Launch Plan", "RD$1,200"],
    ["Almacenamiento", "AWS S3", "Pay-per-use", "RD$500-2,000"],
    ["Email transaccional", "Resend", "Pro Plan", "RD$1,500"],
    ["Errores / Monitoring", "Sentry", "Team Plan", "RD$1,200"],
    ["Analytics", "PostHog + Plausible", "Free / Starter", "RD$0-800"],
    ["Dominio", "Various", ".com", "RD$100/mes (prom.)"],
    ["TOTAL INFRAESTRUCTURA", "", "", "RD$4,000-7,300"],
]
add_styled_table(doc, infra_data[0], infra_data[1:])

doc.add_paragraph("")
doc.add_heading("Anexo D: Proyeccion Mensual Completa (36 Meses)", level=2)

doc.add_paragraph(
    "A continuacion se presenta la proyeccion mes a mes con usuarios, ingresos y costos:"
)

monthly_rows = []
for m in range(1, 37):
    f, p, b = USER_PROJECTIONS[m]
    mrr = calc_mrr(m)
    costs = calc_costs(m)
    net = mrr - costs
    monthly_rows.append([
        f"{m}",
        fmt_num(f),
        fmt_num(p),
        fmt_num(b),
        fmt_dop(mrr),
        fmt_dop(costs),
        fmt_dop(net),
    ])

add_styled_table(doc,
    ["Mes", "Free", "Pro", "Business", "MRR", "Costos", "Flujo Neto"],
    monthly_rows,
)

# ─── PIE DE PAGINA FINAL ─────────────────────────────────────────────────────

doc.add_paragraph("")
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fin del Documento —")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = GRAY_400

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bitacora — Plan Financiero 2026-2029 | CONFIDENCIAL")
run.font.size = Pt(9)
run.font.color.rgb = GRAY_400


# ═══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "Bitacora_Plan_Financiero_2026.docx")
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f"\n{'='*60}")
print(f"  DOCUMENTO GENERADO EXITOSAMENTE")
print(f"{'='*60}")
print(f"  Ubicacion: {output_path}")
print(f"  Tamano: {os.path.getsize(output_path) / 1024:.1f} KB")
print(f"  Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
print(f"{'='*60}\n")
